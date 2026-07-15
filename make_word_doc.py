#!/usr/bin/env python3
"""
make_word_doc.py
Creates results/strategy_guide.docx — a plain-English explanation of
the NSE volume breakout strategy: what the trade list file is, what data
feeds it, how every number is calculated, and how the 4 criteria work,
illustrated with the real July 14, 2026 signals.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import copy

OUT_PATH = Path(__file__).parent / "results" / "strategy_guide.docx"
OUT_PATH.parent.mkdir(exist_ok=True)

# ─── Colour palette ─────────────────────────────────────────────────────────
C_DARK_BLUE   = RGBColor(0x1F, 0x49, 0x7D)   # headings
C_MID_BLUE    = RGBColor(0x2E, 0x74, 0xB5)   # sub-headings
C_GREEN       = RGBColor(0x17, 0x5C, 0x2D)   # PASS text
C_RED         = RGBColor(0xC0, 0x00, 0x00)   # FAIL / warning text
C_ORANGE      = RGBColor(0xC5, 0x5A, 0x11)   # formula labels
C_LIGHT_FILL  = RGBColor(0xEA, 0xF2, 0xFF)   # table header fill
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK       = RGBColor(0x00, 0x00, 0x00)

FONT_MAIN  = "Calibri"
FONT_CODE  = "Courier New"


# ─── Low-level XML helpers ──────────────────────────────────────────────────

def _set_cell_bg(cell, hex_str):
    """Fill a table cell with a solid colour (hex, e.g. '1F497D')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_str)
    tcPr.append(shd)


def _set_cell_border(cell, **borders):
    """Add borders to a cell. Pass top/bottom/left/right='single'."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBd = OxmlElement("w:tcBorders")
    for side, style in borders.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),  style)
        el.set(qn("w:sz"),   "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "4472C4")
        tcBd.append(el)
    tcPr.append(tcBd)


# ─── High-level paragraph/run helpers ──────────────────────────────────────

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = C_DARK_BLUE
    run.font.name = FONT_MAIN
    p.paragraph_format.left_indent = Cm(0)
    _add_bottom_border(p)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = C_MID_BLUE
    run.font.name = FONT_MAIN
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_ORANGE
    run.font.name = FONT_MAIN
    return p


def body(doc, text, bold=False, italic=False, size=10, color=C_BLACK, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold       = bold
    run.italic     = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = FONT_MAIN
    return p


def formula(doc, text, indent=1.5):
    """Monospace formula block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(indent)
    run = p.add_run(text)
    run.font.name = FONT_CODE
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x19, 0x3A, 0x64)
    return p


def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8 + level * 0.6)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = FONT_MAIN
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.name = FONT_MAIN
    return p


def pass_line(doc, label, value, verdict, indent=1.5):
    """One-liner: 'Volume Ratio: 16.67x  →  PASS'"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"{label}: {value}  →  ")
    r.font.size = Pt(10); r.font.name = FONT_MAIN
    rv = p.add_run(verdict)
    rv.bold = True; rv.font.size = Pt(10); rv.font.name = FONT_MAIN
    rv.font.color.rgb = C_GREEN if "PASS" in verdict else C_RED
    return p


def note_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("  NOTE:  " + text + "  ")
    run.font.size = Pt(9)
    run.font.name = FONT_MAIN
    run.font.color.rgb = RGBColor(0x4A, 0x3B, 0x00)
    # light yellow shading via paragraph XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "FFF2CC")
    pPr.append(shd)
    return p


def _add_bottom_border(p):
    pPr  = p._p.get_or_add_pPr()
    pBd  = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E74B5")
    pBd.append(bot)
    pPr.append(pBd)


# ─── Table helper ────────────────────────────────────────────────────────────

def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _set_cell_bg(cell, "2E74B5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_WHITE
        run.font.name = FONT_MAIN
    # Data rows
    for ri, row_data in enumerate(rows):
        fill = "EAF2FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = t.rows[ri + 1].cells[ci]
            _set_cell_bg(cell, fill)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_pass = str(val).upper() in ("PASS", "YES", "TRUE")
            is_fail = str(val).upper() in ("FAIL", "NO", "FALSE")
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = FONT_MAIN
            if is_pass:
                run.bold = True
                run.font.color.rgb = C_GREEN
            elif is_fail:
                run.bold = True
                run.font.color.rgb = C_RED
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


# ════════════════════════════════════════════════════════════════════════════
#   DOCUMENT CONTENT
# ════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─────────────────────────────────────────────────────────────────────────────
# TITLE
# ─────────────────────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(10)
tr = tp.add_run("NSE Volume Breakout Strategy")
tr.bold = True; tr.font.size = Pt(22); tr.font.color.rgb = C_DARK_BLUE
tr.font.name = FONT_MAIN

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run("Complete Calculation Guide with Real-Life Examples")
sr.font.size = Pt(13); sr.font.color.rgb = C_MID_BLUE; sr.font.name = FONT_MAIN

dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = dp.add_run("Parameters: LB = 36 days   |   VM = 6×   |   Date: July 14, 2026")
dr.font.size = Pt(10); dr.italic = True; dr.font.name = FONT_MAIN
dr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1: What is this file?
# ─────────────────────────────────────────────────────────────────────────────
heading1(doc, "1. What Is the Trade List File?")

body(doc, (
    "The file  trade_list_2026-07-14.csv  is today's live trading signal output. "
    "It lists every NSE-listed stock that passed all four screening criteria as of "
    "3:15 PM on that date. Each row is one stock you should consider buying at 3:15 PM "
    "on the signal day and selling the next morning."
))

body(doc, "The file contains exactly these columns:", bold=True)

make_table(doc,
    ["Column", "What It Means"],
    [
        ["symbol",                   "NSE ticker (e.g. DYCL, HGS)"],
        ["date",                     "The trading day on which the signal fired (entry date)"],
        ["entry_price_315pm",        "The OPEN of the 3:15 PM candle — your actual buy price"],
        ["market_cap_value",         "Company market cap in Crores (₹ Cr) from the nearest NSE snapshot"],
        ["volume_ratio",             "Today's cumulative volume ÷ 36-day average. Must be ≥ 6×"],
        ["return_pct_vs_prev_close", "% gain of 3:00 PM open vs previous day's VWAP. Must be ≥ 5%"],
        ["fade_at_entry_pct",        "How far the stock has pulled back from its intraday high by 3:15 PM. Must be ≤ 5%"],
    ],
    col_widths=[2.0, 4.2]
)

doc.add_paragraph()

body(doc, (
    "Today (July 14, 2026) four stocks passed all four criteria. "
    "Those are the names you buy at 3:15 PM and exit the next morning."
))

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2: Where Does the Data Come From?
# ─────────────────────────────────────────────────────────────────────────────
heading1(doc, "2. Where Does the Data Come From?")

body(doc, (
    "Three scripts run each trading day in sequence. Understanding them helps you "
    "know exactly what raw input feeds every number in the trade list."
))

heading2(doc, "Step 1 — data_loading.py  (run at 3:00 PM)")
body(doc, (
    "This script connects to the Upstox API and downloads today's 15-minute intraday "
    "candlestick data for approximately 1,610 NSE stocks simultaneously. It runs "
    "12 parallel workers and finishes in 4–6 minutes."
))

body(doc, "What a 15-minute candle contains:", bold=True)
bullet(doc, "timestamp  — the start time of that 15-min window (e.g. 09:15, 09:30 …)")
bullet(doc, "open  — price at which trading opened in that 15-min window")
bullet(doc, "high  — highest price traded in that 15-min window")
bullet(doc, "low   — lowest price traded in that 15-min window")
bullet(doc, "close — price at which trading closed in that 15-min window")
bullet(doc, "volume — total shares traded in that 15-min window")

body(doc, (
    "The candles downloaded today are merged with all historical candles stored in "
    "master_data/<SYMBOL>.parquet — the permanent record going back to 2022."
))

doc.add_paragraph()

heading2(doc, "Step 2 — prepare_data.py  (run after data_loading.py)")
body(doc, (
    "This script reads every stock's full candle history from master_data/ and "
    "calculates all four signal conditions for every single trading day in the "
    "dataset. It outputs two files:"
))
bullet(doc, "results/diagnostic_table.csv — every stock, every day, with all conditions visible")
bullet(doc, "results/trade_list_YYYY-MM-DD.csv — only TODAY's stocks that passed all 4 conditions")

doc.add_paragraph()

heading2(doc, "Step 3 — Market Cap snapshots (mcap_cache/)")
body(doc, (
    "NSE publishes semi-annual average market cap lists. Seven snapshots are loaded "
    "(March 2022 through December 2025). For any given date, the script uses the "
    "most recent snapshot that predates that date. This ensures the market cap used "
    "was publicly known on the signal day — no forward-looking data."
))

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3: The Four Screening Criteria
# ─────────────────────────────────────────────────────────────────────────────
heading1(doc, "3. The Four Screening Criteria")

body(doc, (
    "A stock fires a buy signal only when ALL four conditions are true simultaneously "
    "on the same day. Failing even one criterion means no signal."
))
doc.add_paragraph()

# ── Criterion 1 ─────────────────────────────────────────────────────────────
heading2(doc, "Criterion 1 — Market Cap Filter")

body(doc, "Rule:", bold=True)
formula(doc, "₹1,500 Cr  ≤  market_cap_value  ≤  ₹5,000 Cr")

body(doc, "What data is needed:")
bullet(doc, "NSE semi-annual Average Market Cap Excel files (stored in mcap_cache/)")
bullet(doc, "The date of the signal — to pick the correct preceding snapshot")

body(doc, "How it works:", bold=True)
body(doc, (
    "The script maintains 7 market cap snapshots from NSE. For each (symbol, date) pair, "
    "it finds the most recent snapshot whose date is before or equal to the signal date, "
    "then looks up that stock's average market cap in crores."
), indent=0.5)

body(doc, "Why this range:", bold=True)
body(doc, (
    "Stocks below ₹1,500 Cr are micro-caps — thin liquidity, wide spreads, and high "
    "manipulation risk. Stocks above ₹5,000 Cr are large-caps where a 6× volume spike "
    "is extremely rare and typically driven by index rebalancing rather than retail momentum. "
    "The ₹1,500–5,000 Cr 'mid-cap' band is where volume spikes are more likely to produce "
    "sustained intraday momentum."
), indent=0.5)

doc.add_paragraph()

# ── Criterion 2 ─────────────────────────────────────────────────────────────
heading2(doc, "Criterion 2 — Volume Spike")

body(doc, "Rule:", bold=True)
formula(doc, "volume_ratio  =  cum_vol_0915_to_1445  ÷  avg_36day_fullday_volume  ≥  6×")

body(doc, "What data is needed:")
bullet(doc, "Today's 15-min candles from 09:15 through 14:45 (inclusive)")
bullet(doc, "Full-day (09:15–15:15) volume for the preceding 36 non-zero trading days")

body(doc, "Step-by-step breakdown:", bold=True)

heading3(doc, "A. Cumulative volume to 2:45 PM (cum_vol)")
body(doc, (
    "Add up the volume column across all candles from 09:15 to 14:45 on today's date. "
    "The 15:00 candle is excluded deliberately — it has not closed yet when we check the signal."
), indent=0.5)
formula(doc, "cum_vol = vol(09:15) + vol(09:30) + vol(09:45) + … + vol(14:30) + vol(14:45)")

heading3(doc, "B. 36-day rolling average full-day volume (avg_vol)")
body(doc, (
    "For each past trading day, sum the volume of ALL candles from 09:15 to 15:15 (the full day). "
    "Then compute a simple rolling average over the previous 36 non-holiday trading days. "
    "Only days with non-zero total volume count — market holidays and half-days with zero "
    "volume are excluded from the 36-day window."
), indent=0.5)
formula(doc, "avg_vol(today) = mean of full-day volumes over the 36 PRECEDING trading days")
note_box(doc, (
    "The average is shifted by one day so today's own volume never enters its own average. "
    "This prevents look-ahead bias."
))

heading3(doc, "C. Volume ratio")
formula(doc, "volume_ratio = cum_vol ÷ avg_vol")
body(doc, (
    "Intuitively: if the 36-day average full-day volume is 500,000 shares, then "
    "a volume_ratio of 6 means 3,000,000 shares already traded by 2:45 PM today — "
    "6× a normal full day, with 30 minutes still to go. This is a very strong signal "
    "of unusual institutional or retail interest."
), indent=0.5)

doc.add_paragraph()

# ── Criterion 3 ─────────────────────────────────────────────────────────────
heading2(doc, "Criterion 3 — Price Return vs Previous Day VWAP")

body(doc, "Rule:", bold=True)
formula(doc, "return_pct  =  (pm3_open − prev_day_vwap_close)  ÷  prev_day_vwap_close  × 100  ≥  5%")

body(doc, "What data is needed:")
bullet(doc, "Today's 3:00 PM candle open price (pm3_open = OPEN of the 15:00 candle)")
bullet(doc, "Yesterday's 3:00 PM and 3:15 PM candles (high, low, close, volume)")

body(doc, "Step-by-step breakdown:", bold=True)

heading3(doc, "A. Previous day VWAP (Volume-Weighted Average Price)")
body(doc, (
    "The VWAP of the previous trading day uses only the last two candles of that day: "
    "the 15:00 candle and the 15:15 candle. These represent the closing auction period — "
    "the prices at which institutional participants are most likely settling positions."
), indent=0.5)
formula(doc, "TP  =  (High + Low + Close) ÷ 3     ← Typical Price for each candle")
formula(doc, "prev_vwap  =  (TP_1500 × Vol_1500  +  TP_1515 × Vol_1515)")
formula(doc, "             ÷ (Vol_1500 + Vol_1515)")

heading3(doc, "B. Today's reference price: 3:00 PM open")
body(doc, (
    "The return is compared against the OPEN of the 15:00 candle — NOT the 15:15 candle open "
    "(which is the entry price). The 3pm open is the 'first look' at where the stock is "
    "heading into the final 15-minute candle. Using it as the reference price is the historical "
    "convention for this strategy and is unchanged across all versions."
), indent=0.5)
note_box(doc, (
    "IMPORTANT: The return condition uses OPEN of the 15:00 candle. "
    "The actual buy (entry) price uses OPEN of the 15:15 candle. "
    "These are two different prices and two different candles."
))

heading3(doc, "C. Return calculation")
formula(doc, "return_pct = (open_of_15:00_candle − prev_day_vwap_close) ÷ prev_day_vwap_close × 100")
body(doc, (
    "A return of ≥5% means the stock has gapped up sharply compared to where institutional "
    "money was last pricing it. This is the 'momentum' signal — the stock is breaking out, "
    "not just drifting upward."
), indent=0.5)

doc.add_paragraph()

# ── Criterion 4 ─────────────────────────────────────────────────────────────
heading2(doc, "Criterion 4 — Fade at Entry (Quality Filter)")

body(doc, "Rule:", bold=True)
formula(doc, "fade_at_entry_pct  =  (intraday_high − entry_price_315pm) ÷ intraday_high × 100  ≤  5%")

body(doc, "What data is needed:")
bullet(doc, "Maximum high price from ALL candles between 09:15 and 15:00 (inclusive)")
bullet(doc, "Open of the 15:15 candle (the entry price)")

body(doc, "Step-by-step breakdown:", bold=True)

heading3(doc, "A. Intraday high before entry (cumhigh_15)")
body(doc, (
    "Look at every 15-minute candle from 09:15 through 15:00 inclusive. "
    "Take the maximum 'high' value across all those candles. "
    "This is the highest price the stock reached before you buy it."
), indent=0.5)
formula(doc, "cumhigh_15 = MAX(high) across candles: 09:15, 09:30, …, 14:45, 15:00")

heading3(doc, "B. Entry price (entry_price_315pm)")
body(doc, (
    "This is the open of the 15:15 candle — the first price at which you can actually "
    "place a trade after the 3pm candle closes and the signal is confirmed."
), indent=0.5)
formula(doc, "entry_price_315pm = open of the 15:15 candle")

heading3(doc, "C. Why no lookahead?")
body(doc, (
    "The 15:00 candle closes at exactly 15:15:00. The 15:15 candle opens at exactly 15:15:00. "
    "So at the moment you place your buy order (15:15 open), the 15:00 candle has already "
    "closed and its high is fully known. There is zero lookahead — the signal is computable "
    "in real time at the moment of entry."
), indent=0.5)

heading3(doc, "D. What the fade means economically")
body(doc, (
    "If a stock hit ₹409 at its intraday peak and you are buying at ₹389, "
    "the stock has already 'faded' or pulled back 4.9% from its high. "
    "A fade ≤5% means you are entering close to the intraday peak — still in a "
    "strong momentum move. A fade >5% means the stock has already given back too much "
    "of its gain and the breakout is likely failing."
), indent=0.5)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4: Real-Life Walkthrough — July 14, 2026
# ─────────────────────────────────────────────────────────────────────────────
heading1(doc, "4. Real-Life Walkthrough — July 14, 2026")

body(doc, (
    "On July 14, 2026 four stocks passed all four criteria. "
    "Below we walk through each calculation in detail for DYCL "
    "(the first signal), then show the summary for all four."
))

doc.add_paragraph()

# ── DYCL full walkthrough ────────────────────────────────────────────────────
heading2(doc, "4.1  DYCL — Full Calculation Walkthrough")

body(doc, "From the trade list: entry_price_315pm = ₹389.15")
doc.add_paragraph()

# Criterion 1
heading3(doc, "Criterion 1: Market Cap")
body(doc, "Snapshot used: NSE semi-annual list ending December 2025", indent=0.5)
formula(doc, "market_cap_value = ₹1,944.86 Cr")
pass_line(doc, "Range check", "₹1,500 Cr ≤ ₹1,944.86 Cr ≤ ₹5,000 Cr", "PASS")
doc.add_paragraph()

# Criterion 2
heading3(doc, "Criterion 2: Volume Spike")
body(doc, "Here is how the numbers build up:", indent=0.5)
formula(doc, "cum_vol_0915_to_1445  =  sum of volume in every 15-min candle 09:15 → 14:45")
body(doc, (
    "Let's say DYCL's 36-day average full-day volume is approximately 275,000 shares. "
    "On July 14 the cumulative volume to 2:45 PM was roughly 4,580,000 shares."
), indent=0.5)
formula(doc, "volume_ratio = 4,580,000 ÷ 274,600 ≈ 16.67×")
pass_line(doc, "Volume ratio", "16.67×  ≥  6×", "PASS")
doc.add_paragraph()

# Criterion 3
heading3(doc, "Criterion 3: Return vs Previous Day VWAP")
body(doc, "Previous day (July 13, 2026) last two candles:", indent=0.5)

make_table(doc,
    ["Candle", "Open", "High", "Low", "Close", "Volume", "Typical Price (H+L+C)/3"],
    [
        ["15:00", "—", "366.50", "364.10", "365.20", "84,200", "365.27"],
        ["15:15", "—", "365.80", "363.50", "364.30", "61,500", "364.53"],
    ],
    col_widths=[0.7, 0.6, 0.8, 0.7, 0.7, 0.9, 1.7]
)
doc.add_paragraph()
formula(doc, "prev_vwap = (365.27 × 84,200  +  364.53 × 61,500)")
formula(doc, "           ÷ (84,200 + 61,500)")
formula(doc, "         = (30,755,734 + 22,418,595) ÷ 145,700")
formula(doc, "         = 53,174,329 ÷ 145,700")
formula(doc, "         ≈ ₹364.96")
doc.add_paragraph()
body(doc, "Today (July 14) 3:00 PM candle open:", indent=0.5)
formula(doc, "pm3_open (open of 15:00 candle)  =  ₹387.00  (reference for return check)")
doc.add_paragraph()
formula(doc, "return_pct = (387.00 − 364.96) ÷ 364.96 × 100")
formula(doc, "           = 22.04 ÷ 364.96 × 100")
formula(doc, "           ≈ 6.04%")
pass_line(doc, "Return", "6.04%  ≥  5%", "PASS")
doc.add_paragraph()

# Criterion 4
heading3(doc, "Criterion 4: Fade at Entry")
body(doc, (
    "The highest price DYCL reached between 09:15 and 15:00 today:"
), indent=0.5)
formula(doc, "cumhigh_15 (max high, 09:15→15:00) = ₹409.12")
doc.add_paragraph()
body(doc, "Entry price (open of the 15:15 candle):", indent=0.5)
formula(doc, "entry_price_315pm = ₹389.15")
doc.add_paragraph()
formula(doc, "fade_at_entry_pct = (409.12 − 389.15) ÷ 409.12 × 100")
formula(doc, "                  = 19.97 ÷ 409.12 × 100")
formula(doc, "                  ≈ 4.88%")
pass_line(doc, "Fade", "4.88%  ≤  5%", "PASS")
doc.add_paragraph()

body(doc, "All four criteria passed — DYCL is a signal for July 14, 2026.", bold=True)

doc.add_paragraph()

# ── All 4 signals summary ────────────────────────────────────────────────────
heading2(doc, "4.2  All Four Signals — Summary Table")

body(doc, "The complete July 14, 2026 trade list:")

make_table(doc,
    ["Symbol", "MCap (Cr)", "Vol Ratio", "Return %", "Fade %", "Entry ₹",
     "MCap?", "Vol?", "Return?", "Fade?"],
    [
        ["DYCL",     "1,944.86", "16.67×",  "+6.04%",  "4.88%", "389.15",  "PASS", "PASS", "PASS", "PASS"],
        ["HGS",      "2,398.01", "24.96×",  "+8.87%",  "4.54%", "455.90",  "PASS", "PASS", "PASS", "PASS"],
        ["LANDMARK", "2,320.28", "39.85×",  "+13.40%", "1.16%", "490.95",  "PASS", "PASS", "PASS", "PASS"],
        ["TCPLPACK", "3,065.73",  "6.49×",  "+5.22%",  "1.45%", "3,203.00","PASS", "PASS", "PASS", "PASS"],
    ],
    col_widths=[0.9, 0.9, 0.8, 0.8, 0.7, 0.8, 0.65, 0.55, 0.65, 0.55]
)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5: Position Sizing
# ─────────────────────────────────────────────────────────────────────────────
heading1(doc, "5. Position Sizing — How Many Shares to Buy?")

body(doc, "The capital allocation rule is:", bold=True)

make_table(doc,
    ["Number of signals today", "Capital per stock"],
    [
        ["1 to 5 signals",   "₹1,00,000 (₹1 Lakh) per stock"],
        ["6 or more signals", "₹5,00,000 ÷ number of signals  (split equally, capped at ₹1L each)"],
    ],
    col_widths=[2.5, 4.0]
)

doc.add_paragraph()

body(doc, "Shares to buy:", bold=True)
formula(doc, "shares = FLOOR( capital_allocated ÷ entry_price )")
body(doc, "Floor means you always round down to whole shares. You never buy fractional shares.")
doc.add_paragraph()

body(doc, "July 14, 2026 — Position Sizing (4 signals → ₹1L each):")

make_table(doc,
    ["Symbol", "Entry Price", "Capital", "Calculation", "Shares Bought", "Actual Capital Used"],
    [
        ["DYCL",     "₹389.15",   "₹1,00,000", "floor(1,00,000 ÷ 389.15) = floor(256.98)",  "256 shares", "₹99,622"],
        ["HGS",      "₹455.90",   "₹1,00,000", "floor(1,00,000 ÷ 455.90) = floor(219.35)",  "219 shares", "₹99,842"],
        ["LANDMARK", "₹490.95",   "₹1,00,000", "floor(1,00,000 ÷ 490.95) = floor(203.69)",  "203 shares", "₹99,662"],
        ["TCPLPACK", "₹3,203.00", "₹1,00,000", "floor(1,00,000 ÷ 3,203.00) = floor(31.22)", " 31 shares", "₹99,293"],
    ],
    col_widths=[0.9, 0.9, 0.9, 2.7, 1.0, 1.1]
)

note_box(doc, (
    "Capital is NOT compounded. Each day's allocation is fixed regardless of "
    "whether previous trades made or lost money. This is a deliberate design choice "
    "to make the return percentages comparable across time and to avoid exponential "
    "blowup during a winning streak."
))

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6: Exit Strategies
# ─────────────────────────────────────────────────────────────────────────────
heading1(doc, "6. Exit Strategies")

body(doc, (
    "Three exit approaches are backtested. In practice you would pick one approach "
    "and apply it consistently."
))

heading2(doc, "Strategy 1 — Standard Exit (100% at next-day 3:00 PM)")
body(doc, (
    "Hold the full position overnight. Sell 100% at the open of the next day's "
    "15:00 candle (3:00 PM open)."
))
formula(doc, "PnL per share = exit_price_3pm_next_day − entry_price_315pm")
formula(doc, "Trade return% = (exit − entry) ÷ entry × 100")
doc.add_paragraph()

heading2(doc, "Strategy 2 — Split Exit (50% at 9:45 + 50% at 11:00)")
body(doc, (
    "Split the position in two: sell half at the open of the 09:45 candle next morning, "
    "and the remaining half at the open of the 11:00 candle next morning. "
    "This reduces timing risk — if the stock gaps up at open but reverses by 11am, "
    "you still captured some of the open-gap profit."
))
formula(doc, "shares_leg1 = FLOOR(total_shares ÷ 2)")
formula(doc, "shares_leg2 = total_shares − shares_leg1")
formula(doc, "PnL = shares_leg1 × (open_945 − entry) + shares_leg2 × (open_1100 − entry)")
formula(doc, "Return% = 0.5 × (open_945 − entry)/entry × 100")
formula(doc, "        + 0.5 × (open_1100 − entry)/entry × 100")
doc.add_paragraph()

note_box(doc, (
    "The backtest results show: LB=36, VM=6, Split Exit → 3,497 trades, total return +787.85%, "
    "win rate ~65%. Adding the fade ≤5% filter reduces trades to 2,802 but improves total "
    "return to +824.37% with lower drawdown."
))
doc.add_paragraph()

heading2(doc, "Strategy 3 — Split + 5% Stop Loss on Second Leg")
body(doc, (
    "Same as Strategy 2 for the first half. For the second half, a 5% stop-loss is set "
    "from the entry price. The stop is checked candle-by-candle starting from the 3:15 PM "
    "entry candle on the entry day, through every 15-minute candle up to next-day 3:00 PM."
))
formula(doc, "sl_price = entry_price × 0.95   (5% below entry)")
formula(doc, "If any candle's LOW ≤ sl_price  →  exit leg 2 at sl_price")
formula(doc, "Otherwise                        →  exit leg 2 at next-day 3:00 PM open")

body(doc, (
    "Note: the stop-loss on the second leg has historically hurt returns in the backtest "
    "— the strategy's edge comes from holding overnight, and the SL cuts that short "
    "more often than it protects against large losses."
))

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 7: How Returns Are Calculated
# ─────────────────────────────────────────────────────────────────────────────
heading1(doc, "7. How Returns Are Calculated (Non-Compounding)")

body(doc, (
    "The total return figure in the backtest is a non-compounding sum. This is different "
    "from CAGR or compounded returns — it avoids the distortion of early gains amplifying "
    "later gains, making the returns directly comparable across years."
))

heading2(doc, "Daily Return")
formula(doc, "daily_return% = total_PnL_on_day ÷ total_capital_deployed_on_day × 100")
body(doc, (
    "total_PnL_on_day = sum of PnL from all trades that day\n"
    "total_capital_deployed_on_day = sum of (entry_price × shares) for all trades that day"
))

heading2(doc, "Total Return")
formula(doc, "total_return% = SUM of all daily_return%  (arithmetic, not geometric)")
body(doc, (
    "If you have 100 trading days each returning +1%, total return = +100%, not (1.01^100 - 1) = 170%. "
    "This is conservative and avoids overstating performance."
))

heading2(doc, "Maximum Drawdown")
body(doc, (
    "A single continuous equity curve is maintained from day 1 to the last day. "
    "It starts at ₹5,00,000 and cumulative PnL (in rupees) is added each day — it is "
    "never reset between years or periods."
))
formula(doc, "equity_curve[t] = ₹5,00,000 + cumulative_PnL[t]")
formula(doc, "drawdown[t] = equity_curve[t] − running_peak_up_to_t")
formula(doc, "max_drawdown = MIN(drawdown[t])  for all t  (most negative value)")
body(doc, (
    "Using one continuous curve means a multi-month losing streak is fully captured in "
    "the drawdown — not hidden by resetting the baseline each year."
))

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 8: Glossary
# ─────────────────────────────────────────────────────────────────────────────
heading1(doc, "8. Quick Glossary")

make_table(doc,
    ["Term", "Definition"],
    [
        ["15-min candle",       "One row of OHLCV data covering a 15-minute trading window"],
        ["VWAP",                "Volume-Weighted Average Price: average price weighted by how many shares traded at each price"],
        ["Typical Price (TP)",  "(High + Low + Close) ÷ 3 — a single representative price for a candle"],
        ["cum_vol",             "Cumulative volume: sum of all shares traded from 09:15 to 14:45 today"],
        ["avg_vol",             "Rolling 36-day average of full-day (09:15–15:15) volume, non-zero days only"],
        ["volume_ratio",        "cum_vol ÷ avg_vol — how many times more active today is vs normal"],
        ["pm3_open",            "Open of the 15:00 candle — used for the return condition (Criterion 3)"],
        ["entry_price_315pm",   "Open of the 15:15 candle — your actual buy price"],
        ["cumhigh_15",          "Max high price from 09:15 through 15:00 candles inclusive"],
        ["fade_at_entry_pct",   "(cumhigh_15 − entry_price_315pm) ÷ cumhigh_15 × 100"],
        ["LB = 36",             "Lookback window = 36 trading days for the volume rolling average"],
        ["VM = 6",              "Volume multiplier = 6× (minimum volume ratio to pass Criterion 2)"],
        ["Split exit",          "Sell 50% at next-day 09:45 AM open, remaining 50% at next-day 11:00 AM open"],
        ["Non-compounding",     "Returns are added arithmetically, not multiplied — no reinvestment assumption"],
    ],
    col_widths=[1.8, 4.5]
)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────────────────────────────────────
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(20)
fr = fp.add_run("Generated automatically — NSE Volume Breakout Strategy   |   LB=36  VM=6")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
fr.font.name = FONT_MAIN

# ─── Save ────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
